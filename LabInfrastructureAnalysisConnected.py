#==================================================================================================================================================================
#   NASA GRC Lab Infrastructure Data Analysis Tool
#   Brandon Hearley - LMS
#   4/15/2024
#
#   PURPOSE: Create a web app (using streamlit) to filter data collected and stored on the NASA
#            GRC Lab Infrastructure and generate a Word Report
#
#==================================================================================================================================================================
# SETUP
# Import the necessary modules to run the app and set paths

# Import Modules
import streamlit as st
import os
import glob
import numpy as np
from pymongo.mongo_client import MongoClient
import dns
import certifi
from datetime import date
import docx
from docx.shared import Pt 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_ORIENT, WD_SECTION
import matplotlib.pyplot as plt
from docx.shared import Inches
import io
from bson.binary import Binary
from PIL import Image

# Set Home Directory
data_path = "/mount/src/labinfrastructure/"

#==================================================================================================================================================================
# GENERAL INFORMATION
# Set the web app general information not edited by the user

# Set the page configuration
st.set_page_config(layout="wide")

# Create the Title
st.title("NASA GRC Laboratory Infrastructure Data Collection Analysis")
st.markdown('Set the filter criteria and generate a Word Document report summarizing the information collected on the GRC Lab Infrastructure. \n\n For all questions, please email brandon.l.hearley@nasa.gov')

#==================================================================================================================================================================
# DATA CONNECTION
# Set up the database connection

# Connect to the Database
@st.cache_resource
def init_connection():
    uri = "mongodb+srv://nasagrc:" + st.secrets['mongo1']['password'] + "@nasagrclabdatatest.hnx1ick.mongodb.net/?retryWrites=true&w=majority&appName=NASAGRCLabDataTest"
    return MongoClient(uri, tlsCAFile=certifi.where())

# Create the Database Client
client = init_connection()

# Send a ping to confirm a successful connection
try:
    client.admin.command('ping')
    print("Pinged your deployment. You successfully connected to MongoDB!")
except Exception as e:
    print(e)

# Read the Database
@st.cache_data(ttl=6000)
def get_data():
    db = client['LabData']
    items = db['LabData'].find()
    items = list(items)  # make hashable for st.cache_data
    return items

# Get All Data in Database
all_data = get_data()

#==================================================================================================================================================================
# DATA FILTERING
# Create options for user to filter the data

#Create Divider for Selection of Divisions or Branches
st.subheader('Select Divisions or Branches')
st.markdown('Only select labs from specific divisions or branches. Use the drop down menu to filter by either division or branch, then check all options you would like in the report.')

# Preallocate List of Divisions and Branches
Div = []
Branch = []

# Loop through the database and find unique branch and division values
for q in range(len(all_data)):
    # Get the branch from the database and extract the division
    branch_q = all_data[q]['Branch']
    div_q = branch_q[0:2]

    # Add the each list if unique
    if div_q not in Div:
        Div.append(div_q)
    if branch_q not in Branch:
        Branch.append(branch_q)

# Sort Alphabetically
Div.sort()
Branch.sort()

# Create Drop down menu for filtering by division or branch
filt_opts = ['','Division', 'Branch']
filt_opt1 = st.selectbox('Filter by:',filt_opts,key='filt_opt1')

# Create array of check box values for each list of divisions and branches
Div_Disp = []
for k in range(len(Div)):
    Div_Disp.append(True)
Branch_Disp = []
for k in range(len(Branch)):
    Branch_Disp.append(True)

# Create checkboxes for filtering
# -- Division Checkboxes
if filt_opt1 == 'Division':
    for j in range(len(Div)):
        Div_Disp[j] = st.checkbox(Div[j], value=False, key='div_' + str(j), label_visibility="visible")
# -- Branch Checkboxes
if filt_opt1 == 'Branch':
    for j in range(len(Branch)):
        Branch_Disp[j] = st.checkbox(Branch[j], value=False, key='div_' + str(j), label_visibility="visible")
st.markdown(' ')

# Utility Function - Convert String Number to Have Comma Separators
def convert_num(key):
    if key in st.session_state:
        raw_val = st.session_state[key]
        if raw_val != None:
            test_val = raw_val.replace(',','')
            flag = 0
            try:
                float(test_val)
            except ValueError:
                if test_val == '':
                    st.session_state[key] = None
                else:
                    flag = 1
                    st.error('Not a valid number')
            if flag == 0:
                if len(test_val)>3:
                    num = ''
                    for k in range(len(test_val)):
                        num = test_val[len(test_val)-1-k] + num
                        chk = k+1
                        if chk%3 == 0 and k!= len(test_val)-1:
                            num = ',' + num
                    st.session_state[key] = num

#Create Divider for Selection of Asset Total Cost based on Condition
st.subheader('Select Labs by Asset or Replacement Costs')
st.markdown('Only select labs whose total asset replacement cost for the specified conditions are within a certain range. Use the two inputs to set the minimum and maximum total value of assets that meet the condition criteria for a lab. Use the checkboxes to select which asset conditions you would like to consider.')

# Create grid for total asset values (min and max)
grid_vals1 = st.columns(2)
with grid_vals1[0]:
    min_asset_cost = st.text_input('Min. Value of Total Assets', value=None, key='min_asset_cost_key', on_change=convert_num('min_asset_cost_key'))
with grid_vals1[1]:
    max_asset_cost = st.text_input('Max. Value of Total Assets', value=None, key='max_asset_cost_key', on_change=convert_num('max_asset_cost_key'))

# Create grid fo asset value checkboxes
grid = st.columns(4)
with grid[0]:
    asset_chk1 = st.checkbox('Poor', value=True)
with grid[1]:
    asset_chk2 = st.checkbox('Fair', value=True)
with grid[2]:
    asset_chk3 = st.checkbox('Good', value=True)
with grid[3]:
    asset_chk4 = st.checkbox('Excellent', value=True)
st.markdown(' ')

#Create Divider for Selection of Total Lab Replacement Cost
st.markdown('Only select labs whose total replacement cost is within a certain range.  Use the two inputs to set the minimum and maximum total value of assets that meet the condition criteria for a lab.')

# Create grid for total replacement cost (min and max)
grid_vals2 = st.columns(2)
with grid_vals2[0]:
    min_tot_cost = st.text_input('Min. Value of Total Replacement Cost', value=None, key='min_tot_cost_key', on_change=convert_num('min_tot_cost_key'))
with grid_vals2[1]:
    max_tot_cost = st.text_input('Max. Value of Total Replacement Cost', value=None, key='max_tot_cost_key', on_change=convert_num('max_tot_cost_key'))
st.markdown(' ')

# Create Drop down menu for selecting entry status
status_opts = ['All Entries (Draft and Final)','Final Only']
status_opt1 = st.selectbox('Select Status of Data Entries:',status_opts,key='status_opt1')
st.markdown(' ')

#==================================================================================================================================================================
# FILTER DATA
# Find all records that meet the selection criteria

# Create the Filter button
st.markdown('Filter all data in the database for the above criteria and write to a report. Once filtered, select the "Download Report" button to download the Word Document.')
if st.button('Filter Data'):
    # Get the Filtering Criteria
    # -- Get List of Divisions
    Div_List = []
    for j in range(len(Div_Disp)):
        if Div_Disp[j] == True:
            Div_List.append(Div[j])

    # -- Get List of Branches
    Branch_List = []
    for j in range(len(Branch_Disp)):
        if Branch_Disp[j] == True:
            Branch_List.append(Branch[j])
    
    # -- Get List of Asset Criteria
    Asset_Cond_List = []
    if asset_chk1 == True:
        Asset_Cond_List.append('Poor')
    if asset_chk2 == True:
        Asset_Cond_List.append('Fair')
    if asset_chk3 == True:
        Asset_Cond_List.append('Good')
    if asset_chk4 == True:
        Asset_Cond_List.append('Excellent')

    # -- Get Asset Value Limits
    if st.session_state['min_asset_cost_key'] != None:
        min_asset_cost = float(st.session_state['min_asset_cost_key'].replace(',',''))
    else:
        min_asset_cost = None
    if st.session_state['max_asset_cost_key'] != None:
        max_asset_cost = float(st.session_state['max_asset_cost_key'].replace(',',''))
    else:
        max_asset_cost = None

    # -- Get Replacement Value Limits
    if st.session_state['min_tot_cost_key'] != None:
        min_tot_cost = float(st.session_state['min_tot_cost_key'].replace(',',''))
    else:
        min_tot_cost = None
    if st.session_state['max_tot_cost_key'] != None:
        max_tot_cost = float(st.session_state['max_tot_cost_key'].replace(',',''))
    else:
        max_tot_cost = None

    # -- Get Status Criteria
    if st.session_state['status_opt1'] == 'All Entries (Draft and Final)':
        status_choice = ['Draft','Final']
    else:
        status_choice = ['Final']

    # -- Create Criteria Dictionary
    criteria= {
             'Div':Div_List,
             'Branches':Branch_List,
             'AssetValCond':Asset_Cond_List,
             'AssetVal':[min_asset_cost, max_asset_cost],
             'RepCost':[min_tot_cost, max_tot_cost],
            'Status':status_choice}

    # Preallocate Dictionary to store records that meet the filtering criteria
    FilesOut = {}
    
    # Evaluate criteria for each record
    for q in range(len(all_data)):
        # Get the individual record
        record = all_data[q]

        # Get the Directorate, Division, and Branch
        Branch = record['Branch']
        Div = Branch[0:2]
        Direc = Branch[0]

        # Set Criteria Flag
        # -- Initialize to 1 (write record to report), if criteria is NOT met, change to 0
        crit_flag = 1 

        # Evaluate Division Criteria
        if criteria['Div'] != None:
            if Div not in criteria['Div']:
                crit_flag = 0

        # Evaluate Branch Criteria
        if criteria['Branches'] != None:
            if Branch not in criteria['Branches']:
                crit_flag = 0

        # Evaluate Total Asset Cost Criteria
        num_assets  = int(record['Number of Assets'])
        tot_asset_cost = 0

        # -- Loop through all assets to get total Asset Cost
        for k in range(num_assets):
            if record['T1-Asset Condition'][k] in criteria['AssetValCond'] and record['T1-Replacement Cost ($)'][k] != None:
                tot_asset_cost = tot_asset_cost + float(record['T1-Replacement Cost ($)'][k])

        # -- Evaluate Asset Value Criteria
        if criteria['AssetVal'][0] != None or criteria['AssetVal'][1] != None:
            if criteria['AssetVal'][0] != None and tot_asset_cost < criteria['AssetVal'][0]:
                crit_flag = 0
            if criteria['AssetVal'][1] != None and tot_asset_cost > criteria['AssetVal'][1]:
                crit_flag = 0
        
        # Evaluate Total Replacement Cost Criteria
        if criteria['RepCost'][0] != None or criteria['RepCost'][1] != None:
            if record['Estimated Cost to Replace Entire Laboratory/Capability ($)'] == None:
                crit_flag = 0
            elif criteria['RepCost'][0] != None and record['Estimated Cost to Replace Entire Laboratory/Capability ($)'] < criteria['RepCost'][0]:
                crit_flag = 0
            elif criteria['RepCost'][1] != None and record['Estimated Cost to Replace Entire Laboratory/Capability ($)'] > criteria['RepCost'][1]:
                crit_flag = 0

        # Evaluate Entry Status Criteria
        if record['Status'] not in criteria['Status']:
            crit_flag= 0

        # Write records that meet filter criteria to the output dictionary
        if crit_flag == 1:
            div_keys = list(FilesOut.keys())
            if Div not in div_keys:
                FilesOut[Div] = {}

            branch_keys = list(FilesOut[Div].keys())
            if Branch not in branch_keys:
                FilesOut[Div][Branch] = []

            FilesOut[Div][Branch].append(q)

#==================================================================================================================================================================
# WRITE THE REPORT
# Write a Word document with filtered data
    
    # Utility Functions
    # -- Change Orientation of the page from landscape to portriat (and vice versa) 
    def change_orientation():
        current_section = doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height
        return new_section

    # Format Entires
    # -- Format the data for excel writing
    def format_values(test_val, key):
        if test_val == None:
            val = ''
        else:
            if key == "string":
                val = str(test_val)
            if key == "money":
                if isinstance(test_val,str) == False:
                    test_val = str(int(test_val))
                val = ''
                for k in range(len(test_val)):
                    val = test_val[len(test_val)-1-k] + val
                    chk = k+1
                    if chk%3 == 0 and k!= len(test_val)-1:
                        val = ',' + val
            if key == "WBS":
                if len(test_val) > 6:
                    val = test_val[0:6]
                else:
                    val = test_val
        return val

    # Create the Document
    doc = docx.Document() 

    # Create the Title Page
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = para.add_run('NASA GRC Lab Infrastructure Data')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(18)
    run1.bold = True

    # Create Date Time Stamp
    today = date.today()
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = para.add_run('Created on: ' + str(today))
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(14)

    # Write Filter Criteria to the title page
    run_lab1 = doc.add_paragraph().add_run('Filer Criteria:')
    run_lab1.font.name = 'Times New Roman'
    run_lab1.font.size = Pt(14)

    # -- Write Divisions
    if st.session_state['filt_opt1'] == 'Division':
        run_lab1 = doc.add_paragraph().add_run('Divisions Selected:')
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

        div_out = ''
        for k in range(len(criteria['Div'])):
            div_out = div_out + criteria['Div'][k] + ', '
        div_out = div_out[:len(div_out)-2]

        run_lab1 = doc.add_paragraph().add_run(div_out)
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

    # -- Write Branches
    if st.session_state['filt_opt1'] == 'Branch':
        run_lab1 = doc.add_paragraph().add_run('Branches Selected:')
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

        branch_out = ''
        for k in range(len(criteria['Branches'])):
            branch_out = branch_out + criteria['Branches'][k] + ', '
        branch_out = branch_out[:len(branch_out)-2]

        run_lab1 = doc.add_paragraph().add_run(branch_out)
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

    # -- Write Asset Costs
    if criteria['AssetVal'][0] != None or criteria['AssetVal'][1]!= None:
        
        run_lab1 = doc.add_paragraph().add_run('Total Asset Value Range:')
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

        cond_out = '    Asset Conditions:'
        for k in range(len(criteria['AssetValCond'])):
            cond_out = cond_out + criteria['AssetValCond'][k] + ', '
        cond_out = cond_out[:len(cond_out)-2]

        run_lab1 = doc.add_paragraph().add_run(cond_out)
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

        val_frmt = format_values(criteria['AssetVal'][0] , "money")
        if val_frmt != '':
            val_Frmt = '$' + val_frmt
        run_lab1 = doc.add_paragraph().add_run('    Minimum Total Asset Value: ' + val_frmt)
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

        val_frmt = format_values(criteria['AssetVal'][1] , "money")
        if val_frmt != '':
            val_Frmt = '$' + val_frmt
        run_lab1 = doc.add_paragraph().add_run('    Maximum Total Asset Value: ' + val_frmt)
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

    # -- Write Total Costs
    if criteria['RepCost'][0] != None or criteria['RepCost'][1]!= None:

        run_lab1 = doc.add_paragraph().add_run('Total Relacement Cost Range:')
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

        val_frmt = format_values(criteria['RepCost'][0] , "money")
        if val_frmt != '':
            val_Frmt = '$' + val_frmt
        run_lab1 = doc.add_paragraph().add_run('    Minimum Total Replacement Cost: ' + val_frmt)
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

        val_frmt = format_values(criteria['RepCost'][1] , "money")
        if val_frmt != '':
            val_Frmt = '$' + val_frmt
        run_lab1 = doc.add_paragraph().add_run('    Maximum Total Replacement Cost: ' + val_frmt)
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(12)

    # -- Status
    # Write Filter header if it doesn't exist
    run_lab1 = doc.add_paragraph().add_run('Entry Status: ' + st.session_state['status_opt1'])
    run_lab1.font.name = 'Times New Roman'
    run_lab1.font.size = Pt(12)

    # Loop Through Divisions
    divisions = list(FilesOut.keys())
    divisions.sort()

    # Set total asset and estimated cost
    tot_cost_all = 0
    tot_asset_cost_all = 0

    # Loop through divisions
    for d in range(len(divisions)):
        # Get list of branches
        branches = list(FilesOut[divisions[d]].keys())
        branches.sort()
        # Loop through the branches
        for b in range(len(branches)):

            # Get List of files
            files = FilesOut[divisions[d]][branches[b]]

            # Loop through individual labs
            for q in range(len(files)):
                # Get the data for the individual record
                record = all_data[files[q]]

                # -- Asset Table
                num_assets = record['Number of Assets']

                # Define the array in the database and format types
                col_dict = {
                    'T1-Asset Name':"string",
                    'T1-Location (Bldg/Rm)':"string",
                    'T1-Age (yrs)':"string",
                    'T1-Acquisition Year':"string",
                    'T1-Expected Year of Obsolescence':"string",
                    'T1-Asset Condition':"string",
                    'T1-Replacement Cost ($)':"money",
                    'T1-Impact to Capability if Lost':"string",
                    'T1-Associated Software':"string",
                    'T1-Inlcudes IT Hardware?':"string",
                    'T1-Replacement':"string"}

                # Get list of array names
                col_keys = list(col_dict.keys())

                # Loop through each asset and write to the table
                for j in range(num_assets):
                    for k in range(len(col_keys)):
                        try:
                            tot_asset_cost_all = tot_asset_cost_all + record['T1-Replacement Cost ($)'][j]
                        except:
                            pass

                # Get Total Estimated Cost
                
                # -- Estimated Cost to Replace Entire Laboratory/Capability ($):
                key = 'Estimated Cost to Replace Entire Laboratory/Capability ($)'
                val = record[key]
                try:
                    tot_cost_all = tot_cost_all + val
                except:
                    pass

                if tot_cost_all < tot_ass_cost_all:
                    sr.write(record['Laboratory/Capability Name'])

    # Write Costs
    val_frmt = format_values(tot_cost_all, "money")
    run_lab1 = doc.add_paragraph().add_run('Total Estimated Lab Cost' + ': ' + val_frmt)
    run_lab1.font.name = 'Times New Roman'
    run_lab1.font.size = Pt(12)

    val_frmt = format_values(tot_asset_cost_all, "money")
    run_lab1 = doc.add_paragraph().add_run('Total Estimated Asset Cost' + ': ' + val_frmt)
    run_lab1.font.name = 'Times New Roman'
    run_lab1.font.size = Pt(12)

    for d in range(len(divisions)):
        # Start on New Page
        doc.add_page_break()

        # Write the Division Name
        run_lab1 = doc.add_paragraph().add_run(divisions[d])
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(18)
        run_lab1.bold = True

        # Get list of branches
        branches = list(FilesOut[divisions[d]].keys())
        branches.sort()

        # Loop through the branches
        for b in range(len(branches)):
            run_lab1 = doc.add_paragraph().add_run(branches[b])
            run_lab1.font.name = 'Times New Roman'
            run_lab1.font.size = Pt(14)
            run_lab1.bold = True

            # Get List of files
            files = FilesOut[divisions[d]][branches[b]]

            # Loop through individual labs
            for q in range(len(files)):
                # Get the data for the individual record
                record = all_data[files[q]]

                # Write the lab name
                run_lab1 = doc.add_paragraph().add_run(record['Laboratory/Capability Name'])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # HEADER: Laboratory/Capability Information
                run_lab1 = doc.add_paragraph().add_run('Laboratory/Capability Information')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Point of Contact
                key = 'Point of Contact'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Branch
                key = 'Branch'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Laboratory/Capability Description
                key = 'Laboratory/Capability Description'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Laboratory/Capability Website
                key = 'Laboratory/Capability Website'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Challenges in sustaining this laboratory/capability
                key = 'Challenges in sustaining this laboratory/capability'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Age (yrs):
                key = 'Age (yrs)'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + str(record[key]))
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Condition:
                key = 'Condition'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Asset Table
                num_assets = record['Number of Assets']
                if num_assets > 0:
                    # Change the page orientation
                    change_orientation()

                    # Write the Asset Table
                    run_lab1 = doc.add_paragraph().add_run('Assets:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)

                    # Create the first row with headers
                    table = doc.add_table(rows=1, cols=11) 
                    row = table.rows[0].cells 
                    row[0].text = 'Asset Name'
                    row[1].text = 'Location (Bldg/Rm)'
                    row[2].text = 'Age (yrs)'
                    row[3].text = 'Asset Date of Entry'
                    row[4].text = 'Expected Date of Obsolescence'
                    row[5].text = 'Asset Condition'
                    row[6].text = 'Replacement Cost ($)'
                    row[7].text = 'Impact to Capability if Lost'
                    row[8].text = 'Associated Software/Required OS'
                    row[9].text = 'IT Hardware Repalcement?'
                    row[10].text = 'Part or Full Replacement?'

                    # Define the array in the database and format types
                    col_dict = {
                        'T1-Asset Name':"string",
                        'T1-Location (Bldg/Rm)':"string",
                        'T1-Age (yrs)':"string",
                        'T1-Acquisition Year':"string",
                        'T1-Expected Year of Obsolescence':"string",
                        'T1-Asset Condition':"string",
                        'T1-Replacement Cost ($)':"money",
                        'T1-Impact to Capability if Lost':"string",
                        'T1-Associated Software':"string",
                        'T1-Inlcudes IT Hardware?':"string",
                        'T1-Replacement':"string"}

                    # Get list of array names
                    col_keys = list(col_dict.keys())

                    # Loop through each asset and write to the table
                    for j in range(num_assets):
                        row = table.add_row().cells
                        for k in range(len(col_keys)):
                            try:
                                val = record[col_keys[k]][j]
                                val_frmt = format_values(val, col_dict[col_keys[k]])
                                row[k].text = val_frmt
                            except:
                                pass
                    table.style = 'Light Grid Accent 4'

                    # Set Font Fize
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'

                    # Set Column Widths
                    table.autofit = False
                    table.allow_autofit =False
                    col_widths =  [0.85,0.79,0.44,0.63,0.94,0.75,0.94,1.19,0.81,1,0.94]
                    for row in table.rows:
                        for k in range(len(col_keys)):
                            row.cells[k].width = Inches(col_widths[k])
                    for k in range(len(col_keys)):
                        table.columns[k].width = Inches(col_widths[k])
                    
                    run_lab1 = doc.add_paragraph().add_run('')

                    # Write Asset Images
                    if 'T7-Asset Image' in list(record.keys()):
                        if len(record['T7-Asset Image']) > 0:
                             # Write the Asset Table
                            run_lab1 = doc.add_paragraph().add_run('Assets Images:')
                            run_lab1.font.name = 'Times New Roman'
                            run_lab1.font.size = Pt(11)
        
                            # Create the first row with headers
                            table = doc.add_table(rows=1, cols=3) 
                            row = table.rows[0].cells 
                            row[0].text = 'Asset Name'
                            row[1].text = 'Image'
                            row[2].text = 'Notes'

                            for j in range(len(record['T7-Asset Image'])):
                                row = table.add_row().cells
                                row[0].text = record['T7-Asset Image Label'][j]
                                row[2].text = record['T7-Asset Image Notes'][j]

                                try:
                                    pil_img = Image.open(io.BytesIO(record['T7-Asset Image'][j]))
                                    plt.imshow(pil_img)
                                    plt.tick_params(left = False, right = False , labelleft = False , labelbottom = False, bottom = False) 
                                    for pos in ['right', 'top', 'bottom', 'left']: 
                                        plt.gca().spines[pos].set_visible(False) 
                                    plt.savefig(os.path.join(data_path,'Asset_Img_' + str(j)+'.png'))
                                    plt.close()
                                    pic = os.path.join(data_path,'Asset_Img_' + str(j)+'.png')  # path image
                                    cell = table.rows[j+1].cells[1]  # position specific of image count=row and 3=column
                                    paragraph = cell.paragraphs[0]
                                    run = paragraph.add_run()
                                    run.add_picture(pic, width=Inches(3), height=Inches(2))  # size image
                                except:
                                    temp=1 # Do Nothing

                        table.style = 'Light Grid Accent 4'
    
                        # Set Font Fize
                        for row in table.rows:
                            for cell in row.cells:
                                paragraphs = cell.paragraphs
                                for paragraph in paragraphs:
                                    for run in paragraph.runs:
                                        font = run.font
                                        font.size= Pt(10)
                                        font.name = 'Times New Roman'

                    # Change page orientation back
                    change_orientation()

                # -- Sustainment Funding Source:
                key = 'Sustainment Funding Source'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Funding Table
                num_fund = record['Number of Funding Sources']
                if num_fund > 0:
                    # Create the Funding Table
                    run_lab1 = doc.add_paragraph().add_run('Funding Sources:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)

                    # Create the first row with ehaders
                    table2 = doc.add_table(rows=1, cols=4) 
                    row = table2.rows[0].cells 
                    row[0].text = 'Funding Source'
                    row[1].text = 'Funding Start Date'
                    row[2].text = 'Funding End Date'
                    row[3].text = 'Funding Amount per Year ($)'

                    # Define the array in the database and format types
                    col_dict = {
                        'T3-Funding Source':"string",
                        'T3-Funding Start Date':"string",
                        'T3-Funding End Date':"string",
                        'T3-Funding Amount per Year ($)':"money"
                               }
                    # Get list of array names
                    col_keys = list(col_dict.keys())

                    # Loop through each row and write to the table
                    for j in range(num_fund):
                        row = table2.add_row().cells
                        for k in range(len(col_keys)):
                            val = record[col_keys[k]][j]
                            val_frmt = format_values(val, col_dict[col_keys[k]])
                            row[k].text = val_frmt
                    table2.style = 'Light Grid Accent 4'

                    # Set Font Fize
                    for row in table2.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'
                                    
                    run_lab1 = doc.add_paragraph().add_run('')

                # HEADER: Current Mission/Project Utilization
                run_lab1 = doc.add_paragraph().add_run('Current Mission/Project Utilization')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Project Table
                num_proj = record['Number of Projects']
                if num_proj > 0:
                    # Create hte Project TAble
                    run_lab1 = doc.add_paragraph().add_run('Projects:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)

                    # Create the first row with headers
                    table3 = doc.add_table(rows=1, cols=5) 
                    row = table3.rows[0].cells 
                    row[0].text = 'Mission/Project Name'
                    row[1].text = 'WBS Number'
                    row[2].text = 'Project Use (%)'
                    row[3].text = 'Risk to Project'
                    row[4].text = 'Impact if Laboratory/Capability is Lost'

                    # Define the array in the database and format types
                    col_dict = {
                        'T4-Mission/Project Name':"string",
                        'T4-WBS Number':"WBS",
                        'T4-Project Use (%)':"string",
                        'T4-Risk to Project':"string",
                        'T4-Impact if Laboratory/Capability is Lost':"string"
                    }

                    # Get list of array names
                    col_keys = list(col_dict.keys())

                    # Loop through each row and write to the table
                    for j in range(num_proj):
                        row = table3.add_row().cells
                        for k in range(len(col_keys)):
                            val = record[col_keys[k]][j]
                            val_frmt = format_values(val, col_dict[col_keys[k]])
                            row[k].text = val_frmt
                    table3.style = 'Light Grid Accent 4'

                    # Set Font Fize
                    for row in table3.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'

                    # Create Pie Chart
                    labels = []
                    vals = []
                    for j in range(num_proj):
                        try: 
                            float(record[col_keys[2]][j])
                            labels.append(record[col_keys[0]][j])
                            vals.append(float(record[col_keys[2]][j]))
                        except:
                            temp=1

                    if len(vals) > 0 and max(vals) > 0:
                        run_lab1 = doc.add_paragraph().add_run('')
                        fig, ax = plt.subplots()
                        ax.pie(vals, labels=labels, autopct='%1.0f%%')
                        plt.savefig(os.path.join(data_path,'Project_chart_' + str(q)+'.png'))
                        plt.close()
                        doc.add_picture(os.path.join(data_path,'Project_chart_' + str(q)+'.png'), width=Inches(4), height=Inches(3))

                # HEADER: Utilization History/Impact
                run_lab1 = doc.add_paragraph().add_run('Utilization History/Impact')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- History of capability utilization
                key = 'History of capability utilization'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Major impact and contributions this capability has made possible:
                key = 'Major impact and contributions this capability has made possible'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Overall impact of laboratory/capability is lost:
                key = 'Overall impact of laboratory/capability is lost'
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + record[key])
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # HEADER: History of Down Time Due to Maintenance or Failure
                run_lab1 = doc.add_paragraph().add_run('History of Down Time Due to Maintenance or Failure')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Read Down Time Table
                num_dt = record['Number of Failures']
                if num_dt > 0:
                    # Create the Down Time Table
                    run_lab1 = doc.add_paragraph().add_run('Previous Laboratory/Asset Failures:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)

                    # Create the first row with headers
                    table4 = doc.add_table(rows=1, cols=6) 
                    row = table4.rows[0].cells 
                    row[0].text = 'Asset'
                    row[1].text = 'Start Date'
                    row[2].text = 'Time Down'
                    row[3].text = 'Time Down Unit'
                    row[4].text = 'Additional Notes'
                    row[5].text = 'Impact on Mission/Project'

                    # Define the array in the database and format types
                    col_dict = { 
                        'T5-Asset':"string",
                        'T5-Start Date':"string",
                        'T5-Time Down':"string",
                        'T5-Unit':"string",
                        'T5-Additional Notes':"string",
                        'T5-Impact':"string"
                               }

                    # Get list of array names
                    col_keys = list(col_dict.keys())

                    # Loop through each row and write to the table
                    for j in range(num_dt):
                        row = table4.add_row().cells
                        for k in range(len(col_keys)):
                            val = record[col_keys[k]][j]
                            val_frmt = format_values(val, col_dict[col_keys[k]])
                            row[k].text = val_frmt
                    table4.style = 'Light Grid Accent 4'
                    
                    # Set Font Fize
                    for row in table4.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'
                    run_lab1 = doc.add_paragraph().add_run('')

                # HEADER: Cost
                run_lab1 = doc.add_paragraph().add_run('Cost')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Estimated Cost to Replace Entire Laboratory/Capability ($):
                key = 'Estimated Cost to Replace Entire Laboratory/Capability ($)'
                val = record[key]
                val_frmt = format_values(val, "money")
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + val_frmt)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)
    
        
                # -- Cost of Service Contracts ($):
                key = 'Cost of Service Contracts ($)'
                val = record[key]
                val_frmt = format_values(val, "money")
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + val_frmt)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)
        
                # -- Annual Cost to Operate and Sustain the Lab ($/yr):
                key = 'Annual Cost to Operate and Sustain the Lab ($/yr)'
                val = record[key]
                val_frmt = format_values(val, "money")
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + val_frmt)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)
        
                # -- Cost of Service Contracts ($):
                key = 'Incurred Cost For Downtime ($/yr)'
                val = record[key]
                val_frmt = format_values(val, "money")
                run_lab1 = doc.add_paragraph().add_run(key + ': ' + val_frmt)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Read Divisons Table
                num_div  = record['Number of Divisions (Labor Costs)']
                if num_div > 0:
                    run_lab1 = doc.add_paragraph().add_run('Directorate Labor Division:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)
                    table = doc.add_table(rows=1, cols=2) 
                    row = table.rows[0].cells 
                    row[0].text = 'Directorate'
                    row[1].text = 'Labor Division (%)'

                    # Define the array in the database and format types
                    col_dict = {
                        'T6-Directorate':"string",
                        'T6-Labor Cost (%)':"string"
                               }

                    # Get list of array names
                    col_keys = list(col_dict.keys())

                    # Loop through each row and write to the table
                    for j in range(num_div):
                        row = table.add_row().cells
                        for k in range(len(col_keys)):
                            val = record[col_keys[k]][j]
                            val_frmt = format_values(val, col_dict[col_keys[k]])
                            row[k].text = val_frmt
                    table.style = 'Light Grid Accent 4'

                    # Set Font Fize
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'
                                    
                    run_lab1 = doc.add_paragraph().add_run('')

                    # Create Pie Chart
                    labels = []
                    vals = []
                    for j in range(num_div):
                        labels.append(record[col_keys[0]][j])
                        vals.append(float(record[col_keys[1]][j]))

                    run_lab1 = doc.add_paragraph().add_run('')
                    fig, ax = plt.subplots()
                    ax.pie(vals, labels=labels, autopct='%1.0f%%')
                    plt.savefig(os.path.join(data_path,'Labor_chart_' + str(q)+'.png'))
                    plt.close()
                    doc.add_picture(os.path.join(data_path,'Labor_chart_' + str(q)+'.png'), width=Inches(4), height=Inches(3))

                # Check for Lab Images and write
                if 'Lab Images' in list(record.keys()):
                    run_lab1 = doc.add_paragraph().add_run('Laboratory Images')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(12)
                    run_lab1.bold = True
                    for imcnt in range(len(record['Lab Images'])):
                        pil_img = Image.open(io.BytesIO(record['Lab Images'][imcnt]))
                        plt.imshow(pil_img)
                        plt.tick_params(left = False, right = False , labelleft = False , labelbottom = False, bottom = False) 
                        for pos in ['right', 'top', 'bottom', 'left']: 
                            plt.gca().spines[pos].set_visible(False) 
                        plt.savefig(os.path.join(data_path,'Lab_Img_' + str(imcnt)+'.png'))
                        plt.close()
                        doc.add_picture(os.path.join(data_path,'Lab_Img_' + str(imcnt)+'.png'), width=Inches(4), height=Inches(3))

                # Start on New Page
                doc.add_page_break()

    # Save the Document
    doc.save('Lab Data Output.docx')
    doc_download = doc

    # Write to io and create download button
    bio = io.BytesIO()
    doc_download.save(bio)
    if doc_download:
        st.download_button(
            label="Download Report",
            data=bio.getvalue(),
            file_name='Lab Data Output.docx',
            mime="docx"
        )
