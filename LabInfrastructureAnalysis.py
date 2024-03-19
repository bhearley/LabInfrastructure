# Import Modules
import streamlit as st
import os
import glob

# Set Paths
home = os.getcwd()
data_path = "/mount/src/labinfrastructure/Final/"

# Set the page configuration
st.set_page_config(layout="wide")

# Create the Title
st.title("NASA GRC Laboratory Infrastructure Data Collection Analysis")
st.markdown('Set the filter criteria and generate a Word Document report summarizing the information collected on the GRC Lab Infrastructure. \n\n For all questions, please email brandon.l.hearley@nasa.gov')

#Create Divider for Name and Description
st.subheader('Set Filter Criteria')
st.markdown('Only select labs from specific divisions or branches (check all that apply).')

# Get List of Divisions and Branches
os.chdir(data_path)
files_all = glob.glob('*.txt')

Div = []
Branch = []

for q in range(len(files_all)):
    # Read the Text File
    with open(os.path.join(data_path,files_all[q])) as f:
        lines = f.readlines()

    # Get the Branch Name
    key = 'Branch:'
    for i in range(len(lines)):
        if key in lines[i]:
            val  = lines[i][len(key)+1:len(lines[i])-1]

    if len(val) == 3:
        Direc = val[0]
        if val[0:2] not in Div:
            Div.append(val[0:2])
        if val[0:3] not in Branch:
            Branch.append(val[0:3])
    
Div.sort()
Branch.sort()


# Filter Criteria - By Branch or Division
filt_opts = ['','Division', 'Branch']
filt_opt1 = st.selectbox('Filter by:',filt_opts,key='filt_opt1')

Div_Disp = []
for k in range(len(Div)):
    Div_Disp.append(True)
Branch_Disp = []
for k in range(len(Branch)):
    Branch_Disp.append(True)

if filt_opt1 == 'Division':
    for j in range(len(Div)):
        Div_Disp[j] = st.checkbox(Div[j], value=True, key='div_' + str(j), label_visibility="visible")
if filt_opt1 == 'Branch':
    for j in range(len(Branch)):
        Branch_Disp[j] = st.checkbox(Branch[j], value=True, key='div_' + str(j), label_visibility="visible")

st.markdown(' ')
st.markdown('Only select labs whose total asset replacement cost for the specified conditions are within a certain range. Use the checkboxes to only add an assets value to the total for the lab if it meets that condition.')

grid_vals1 = st.columns(2)
with grid_vals1[0]:
    min_asset_cost = st.number_input('Min. Value of Total Assets', min_value=None, max_value=None, value=0)
with grid_vals1[1]:
    max_asset_cost = st.number_input('Max. Value of Total Assets', min_value=None, max_value=None, value=100000000)

grid = st.columns(4)
with grid[0]:
    asset_chk1 = st.checkbox('Poor', value=True)
with grid[1]:
    asset_chk2 = st.checkbox('Fair', value=True)
with grid[2]:
    asset_chk3 = st.checkbox('Good', value=True)
with grid[3]:
    asset_chk4 = st.checkbox('Excellent', value=True)


def convert_num(key):
    if key in st.session_state:
        test_val = st.session_state[key]
        if test_val.isnumeric() == True:
            if len(test_val)>3:
                num = ''
                for k in range(len(test_val)):
                    num = test_val[len(test_val)-1-k] + num
                    chk = k+1
                    if chk%3 == 0 and k!= len(test_val)-1:
                        num = ',' + num
                st.session_state[key] = num
            

test_val = st.text_input('Test', value="1,000,000", key='test_txt', on_change=convert_num('test_txt'))


st.markdown(' ')
st.markdown('Only select labs whose total replacement cost is within a certain range.')
grid_vals2 = st.columns(2)
with grid_vals2[0]:
    min_tot_cost = st.number_input('Min. Value of Total Replacement Cost', min_value=None, max_value=None, value=0)
with grid_vals2[1]:
    max_tot_cost = st.number_input('Max. Value of Total Replacement Cost', min_value=None, max_value=None, value=100000000)

st.markdown(' ')
st.markdown('Filter all data in the database for the above criteria and write to a report. Once filtered, select the "Download Report" button to download the Word Document.')
if st.button('Filter Data'):
    # Get List of Divisions
    Div_List = []
    for j in range(len(Div_Disp)):
        if Div_Disp[j] == True:
            Div_List.append(Div[j])

    # Get List of Divisions
    Branch_List = []
    for j in range(len(Branch_Disp)):
        if Branch_Disp[j] == True:
            Branch_List.append(Branch[j])
    
    # Get List of Asset Criteria
    Asset_Cond_List = []
    if asset_chk1 == True:
        Asset_Cond_List.append('Poor')
    if asset_chk2 == True:
        Asset_Cond_List.append('Fair')
    if asset_chk3 == True:
        Asset_Cond_List.append('Good')
    if asset_chk4 == True:
        Asset_Cond_List.append('Excellent')

    # Create Criteria Dictionary
    criteria= {
             'Div':Div_List,
             'Branches':Branch_List,
             'AssetValCond':Asset_Cond_List,
             'AssetVal':[min_asset_cost, max_asset_cost],
             'RepCost':[min_tot_cost, max_tot_cost]}

    FilesOut = {}
    
    # Get Organized List of Records
    for q in range(len(files_all)):
        # Read the Text File
        with open(os.path.join(data_path,files_all[q])) as f:
            lines = f.readlines()

        # Get the Branch Name
        key = 'Branch:'
        for i in range(len(lines)):
            if key in lines[i]:
                val  = lines[i][len(key)+1:len(lines[i])-1]

        if len(val) == 3:
            Direc = val[0]
            Div = val[0:2]
            Branch = val[0:3]

        # Set Flag
        crit_flag = 1

        # Evaluate Division Criteria
        if criteria['Div'] != None:
            if Div not in criteria['Div']:
                crit_flag = 0

        # Evaluate Division Criteria
        if criteria['Branches'] != None:
            if Branch not in criteria['Branches']:
                crit_flag = 0

        # Get Total Asset Cost and Filter

        key = 'Number of Assets:'
        for i in range(len(lines)):
            if key in lines[i]:
                val  = lines[i][len(key)+1:len(lines[i])-1]
                line_num = i
        num_assets  = int(val)
        
        data = ''
        for k in range(line_num+2,line_num+2+num_assets):
            data = data + lines[k]
        data= data.split('\n')
        data_all = []
        for k in range(num_assets):
            data_line = data[k]
            data_line = data_line.split('\t')
            data_all.append(data_line)
        tot_asset_cost = 0
        for k in range(num_assets):
            if data_all[k][5] in criteria['AssetValCond']:
                tot_asset_cost = tot_asset_cost + float(data_all[k][6])

        if criteria['AssetVal'][0] != None or criteria['AssetVal'][0] != None:
            if criteria['AssetVal'][0] != None and tot_asset_cost < criteria['AssetVal'][0]:
                crit_flag = 0
            if criteria['AssetVal'][1] != None and tot_asset_cost > criteria['AssetVal'][1]:
                crit_flag = 0
        
        # -- Estimated Cost to Replace Entire Laboratory/Capability ($):
        key = 'Estimated Cost to Replace Entire Laboratory/Capability ($):'
        for i in range(len(lines)):
            if key in lines[i]:
                val  = lines[i][len(key)+1:len(lines[i])-1]
        if criteria['RepCost'][0] != None or criteria['RepCost'][0] != None:
            if criteria['RepCost'][0] != None and tot_asset_cost < criteria['RepCost'][0]:
                crit_flag = 0
            if criteria['RepCost'][1] != None and tot_asset_cost > criteria['RepCost'][1]:
                crit_flag = 0


        if crit_flag == 1:
            div_keys = list(FilesOut.keys())
            if Div not in div_keys:
                FilesOut[Div] = {}

            branch_keys = list(FilesOut[Div].keys())
            if Branch not in branch_keys:
                FilesOut[Div][Branch] = []

            FilesOut[Div][Branch].append(files_all[q])

    # Write the Report
    import docx
    from docx.shared import Pt 
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_BREAK
    from docx.enum.section import WD_ORIENT, WD_SECTION
    import matplotlib.pyplot as plt
    from docx.shared import Inches

    # Utility Function
    def change_orientation():
        current_section = doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height

        return new_section

    # Create the Document
    doc = docx.Document() 

    # Create the Title Page
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = para.add_run('NASA GRC Lab Infrastructure Data')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(18)
    run1.bold = True

    # Loop Through Divisions
    divisions = list(FilesOut.keys())
    divisions.sort()

    for d in range(len(divisions)):
        # Start on New Page
        doc.add_page_break()

        run_lab1 = doc.add_paragraph().add_run(divisions[d])
        run_lab1.font.name = 'Times New Roman'
        run_lab1.font.size = Pt(18)
        run_lab1.bold = True

        # Get list of branches
        branches = list(FilesOut[divisions[d]].keys())
        branches.sort()

        for b in range(len(branches)):
            run_lab1 = doc.add_paragraph().add_run(branches[b])
            run_lab1.font.name = 'Times New Roman'
            run_lab1.font.size = Pt(14)
            run_lab1.bold = True

            # Get List of files
            files = FilesOut[divisions[d]][branches[b]]
            files.sort()

            for q in range(len(files)):
                
                # Read the Text File
                with open(os.path.join(data_path,files[q])) as f:
                    lines = f.readlines()

                
                # -- Laboratory/Capability Name
                key = 'Laboratory/Capability Name:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]

                run_lab1 = doc.add_paragraph().add_run(val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # HEADER: Laboratory/Capability Information
                run_lab1 = doc.add_paragraph().add_run('Laboratory/Capability Information')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Point of Contact
                key = 'Point of Contact:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Point of Contact
                key = 'Branch:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Laboratory/Capability Description
                key = 'Laboratory/Capability Description:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Laboratory/Capability Website
                key = 'Laboratory/Capability Website:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Challenges in sustaining this laboratory/capability
                key = 'Challenges in sustaining this laboratory/capability:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Age (yrs):
                key = 'Age (yrs):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Condition:
                key = 'Condition:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

    
                # -- Asset Table
                key = 'Number of Assets:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_assets  = int(val)
        
                data = ''
                for k in range(line_num+2,line_num+2+num_assets):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_assets):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)

                if num_assets > 0:
                    change_orientation()

                    run_lab1 = doc.add_paragraph().add_run('Assets:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)

                    table = doc.add_table(rows=1, cols=11) 
                    row = table.rows[0].cells 
                    row[0].text = 'Asset Name'
                    row[1].text = 'Location (Bldg/Rm)'
                    row[2].text = 'Age (yrs)'
                    row[3].text = 'Asset Date of Entry'
                    row[4].text = 'Expected Date of Obsolescence'
                    row[5].text = 'Asset Condition'
                    row[6].text = 'Replacement Cost ($)'
                    row[7].text = 'Impact to Capability if Lost'
                    row[8].text = 'Associated Software/Required OS'
                    row[9].text = 'IT Hardware Repalcement?'
                    row[10].text = 'Part or Full Replacement?'

                    for j in range(len(data_all)):
                        row = table.add_row().cells
                        for k in range(11):
                            row[k].text = data_all[j][k]
                    table.style = 'Light Grid Accent 4'

                    # Set Font Fize
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'

                    # Set Column Widths
                    table.autofit = False
                    table.allow_autofit =False
                    col_widths =  [0.85,0.79,0.44,0.63,0.94,0.75,0.94,1.19,0.81,1,0.94]
                    for row in table.rows:
                        for k in range(11):
                            row.cells[k].width = Inches(col_widths[k])
                    for k in range(11):
                        table.columns[k].width = Inches(col_widths[k])
                    
                    run_lab1 = doc.add_paragraph().add_run('')

                    change_orientation()

                # -- Sustainment Funding Source:
                key = 'Sustainment Funding Source:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Funding Table
                key = 'Number of Funding Sources:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_fund  = int(val)

                data = ''
                for k in range(line_num+2,line_num+2+num_fund):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_fund):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)
    
                if num_fund > 0:
                    run_lab1 = doc.add_paragraph().add_run('Funding Sources:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)
                    table2 = doc.add_table(rows=1, cols=4) 
                    row = table2.rows[0].cells 
                    row[0].text = 'Funding Source'
                    row[1].text = 'Funding Start Date'
                    row[2].text = 'Funding End Date'
                    row[3].text = 'Funding Amount per Year ($)'


                    for j in range(len(data_all)):
                        row = table2.add_row().cells
                        for k in range(4):
                            row[k].text = data_all[j][k]
                    table2.style = 'Light Grid Accent 4'

                    # Set Font Fize
                    for row in table2.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'
                                    
                    run_lab1 = doc.add_paragraph().add_run('')

                # HEADER: Current Mission/Project Utilization
                run_lab1 = doc.add_paragraph().add_run('Current Mission/Project Utilization')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # Project Table
                key = 'Number of Projects:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_proj  = int(val)

                data = ''
                for k in range(line_num+2,line_num+2+num_proj):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_proj):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)

                if num_proj > 0:
                    run_lab1 = doc.add_paragraph().add_run('Projects:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)
                    table3 = doc.add_table(rows=1, cols=5) 
                    row = table3.rows[0].cells 
                    row[0].text = 'Mission/Project Name'
                    row[1].text = 'WBS Number'
                    row[2].text = 'Project Use (%)'
                    row[3].text = 'Risk to Project'
                    row[4].text = 'Impact if Laboratory/Capability is Lost'


                    for j in range(len(data_all)):
                        row = table3.add_row().cells
                        for k in range(5):
                            if k == 1:
                                row[k].text = data_all[j][k][0:6]
                            else:
                                row[k].text = data_all[j][k]
                    table3.style = 'Light Grid Accent 4'

                    # Set Font Fize
                    for row in table3.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'

                    # Create Pie Chart
                    labels = []
                    vals = []
                    for j in range(len(data_all)):
                        labels.append(data_all[j][0])
                        vals.append(float(data_all[j][2]))

                    run_lab1 = doc.add_paragraph().add_run('')
                    fig, ax = plt.subplots()
                    ax.pie(vals, labels=labels, autopct='%1.0f%%')
                    plt.savefig(os.path.join(data_path,'Project_chart_' + str(q)+'.png'))
                    doc.add_picture(os.path.join(data_path,'Project_chart_' + str(q)+'.png'), width=Inches(4), height=Inches(3))

                # HEADER: Utilization History/Impact
                run_lab1 = doc.add_paragraph().add_run('Utilization History/Impact')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- History of capability utilization
                key = 'History of capability utilization:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]

                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Major impact and contributions this capability has made possible:
                key = 'Major impact and contributions this capability has made possible:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]

                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # HEADER: History of Down Time Due to Maintenance or Failure
                run_lab1 = doc.add_paragraph().add_run('History of Down Time Due to Maintenance or Failure')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Read Down Time Table
                key = 'Number of Failures:'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_dt  = int(val)
        
                data = ''
                for k in range(line_num+2,line_num+2+num_dt):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_dt):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)

                if num_dt > 0:
                    run_lab1 = doc.add_paragraph().add_run('Previous Laboratory/Asset Failures:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)
                    table4 = doc.add_table(rows=1, cols=5) 
                    row = table4.rows[0].cells 
                    row[0].text = 'Asset'
                    row[1].text = 'Start Date'
                    row[2].text = 'Time Down'
                    row[3].text = 'Time Down Unit'
                    row[4].text = 'Additional Notes'

                    for j in range(len(data_all)):
                        row = table4.add_row().cells
                        for k in range(5):
                            row[k].text = data_all[j][k]
                    table4.style = 'Light Grid Accent 4'
                    
                    # Set Font Fize
                    for row in table4.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'
                    run_lab1 = doc.add_paragraph().add_run('')

                # HEADER: Cost
                run_lab1 = doc.add_paragraph().add_run('Cost')
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(12)
                run_lab1.bold = True

                # -- Estimated Cost to Replace Entire Laboratory/Capability ($):
                key = 'Estimated Cost to Replace Entire Laboratory/Capability ($):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)
    
        
                # -- Cost of Service Contracts ($):
                key = 'Cost of Service Contracts ($):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)
        
                # -- Annual Cost to Operate and Sustain the Lab ($/yr):
                key = 'Annual Cost to Operate and Sustain the Lab ($/yr):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)
        
                # -- Cost of Service Contracts ($):
                key = 'Incurred Cost For Downtime ($/yr):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                run_lab1 = doc.add_paragraph().add_run(key + ' ' + val)
                run_lab1.font.name = 'Times New Roman'
                run_lab1.font.size = Pt(11)

                # -- Read Divisons Table
                key = 'Number of Divisions (Labor Costs):'
                for i in range(len(lines)):
                    if key in lines[i]:
                        val  = lines[i][len(key)+1:len(lines[i])-1]
                        line_num = i
                num_div  = int(val)

        
                data = ''
                for k in range(line_num+2,line_num+2+num_div):
                    data = data + lines[k]
                data= data.split('\n')
                data_all = []
                for k in range(num_div):
                    data_line = data[k]
                    data_line = data_line.split('\t')
                    data_all.append(data_line)

                if num_div > 0:
                    run_lab1 = doc.add_paragraph().add_run('Directorate Labor Division:')
                    run_lab1.font.name = 'Times New Roman'
                    run_lab1.font.size = Pt(11)
                    table = doc.add_table(rows=1, cols=2) 
                    row = table.rows[0].cells 
                    row[0].text = 'Directorate'
                    row[1].text = 'Labor Division (%)'
                    for j in range(len(data_all)):
                        row = table.add_row().cells
                        for k in range(2):
                            row[k].text = data_all[j][k]
                    table.style = 'Light Grid Accent 4'

                    # Set Font Fize
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= Pt(10)
                                    font.name = 'Times New Roman'
                                    
                    run_lab1 = doc.add_paragraph().add_run('')

                    # Create Pie Chart
                    labels = []
                    vals = []
                    for j in range(len(data_all)):
                        labels.append(data_all[j][0])
                        vals.append(float(data_all[j][1]))

                    run_lab1 = doc.add_paragraph().add_run('')
                    fig, ax = plt.subplots()
                    ax.pie(vals, labels=labels, autopct='%1.0f%%')
                    plt.savefig(os.path.join(data_path,'Labor_chart_' + str(q)+'.png'))
                    doc.add_picture(os.path.join(data_path,'Labor_chart_' + str(q)+'.png'), width=Inches(4), height=Inches(3))

                # Start on New Page
                doc.add_page_break()




    # Save the Document
    doc.save('Lab Data Output.docx')

    import io
    doc_download = doc

    bio = io.BytesIO()
    doc_download.save(bio)
    if doc_download:
        st.download_button(
            label="Download Report",
            data=bio.getvalue(),
            file_name='Lab Data Output.docx',
            mime="docx"
        )


